from docx import Document
from docx.shared import Pt

print('Creating new docx document')
document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(60)

document.add_paragraph("0123456789").add_run()

print('Saving document')
document.save('test.docx')
